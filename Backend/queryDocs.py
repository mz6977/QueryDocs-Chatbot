import os
from openai import AzureOpenAI
from dotenv import load_dotenv
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings
from langchain_community.vectorstores.azuresearch import AzureSearch
from langchain_core.documents import Document
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import HttpResponseError


load_dotenv()


AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_CHAT_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME")
AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_ENDPOINT")
AZURE_SEARCH_ADMIN_KEY = os.getenv("AZURE_SEARCH_ADMIN_KEY")
AZURE_SEARCH_INDEX_NAME = os.getenv("AZURE_SEARCH_INDEX_NAME")
 

if not all([AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_CHAT_DEPLOYMENT_NAME,
            AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME, AZURE_OPENAI_API_VERSION, AZURE_SEARCH_ENDPOINT, 
            AZURE_SEARCH_ADMIN_KEY, AZURE_SEARCH_INDEX_NAME]):
    raise ValueError("Azure OpenAI environment variables not set. Please check your .env file.")
 

chat_client = AzureOpenAI(
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT
)
 
embeddings_model = AzureOpenAIEmbeddings(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME,
    openai_api_key=AZURE_OPENAI_API_KEY,
    openai_api_version=AZURE_OPENAI_API_VERSION,
    chunk_size=1024
)


vector_store = AzureSearch(
    azure_search_endpoint=AZURE_SEARCH_ENDPOINT,
    azure_search_key=AZURE_SEARCH_ADMIN_KEY,
    index_name=AZURE_SEARCH_INDEX_NAME,
    embedding_function=embeddings_model.embed_query
)


client = SearchClient(
            endpoint=AZURE_SEARCH_ENDPOINT,
            index_name=AZURE_SEARCH_INDEX_NAME,
            credential=AzureKeyCredential(AZURE_SEARCH_ADMIN_KEY)
        )


def extract_text_and_file_extension(file, file_extension):
    
    if file_extension == '.pdf':
        reader = PyPDF2.PdfReader(file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    
    if file_extension == '.docx':
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs]).strip()


def delete_all_documents():
    try:
        # Retrieve all documents
        results = client.search("")  # Empty query retrieves all documents

        # Check for the correct key field in the documents
        documents_to_delete = []
        for doc in results:
            # Replace "id" with your actual key field name if different
            document_key = doc.get("id") or doc.get("@search.documentkey")
            if document_key:
                documents_to_delete.append({"@search.action": "delete", "id": document_key})

        print(documents_to_delete)
        if documents_to_delete:
            # Delete documents
            client.upload_documents(documents=documents_to_delete)
            print("All documents deleted successfully.")
        else:
            print("No documents to delete or no valid keys found.")
    except Exception as e:
        print(f"An error occurred while deleting documents: {e}")


def process_document_for_rag(filename, file_extension, extracted_text):
    try:
        # Chunk the extracted text
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=150,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_text(extracted_text)

        # Prepare documents for vector store with metadata
        docs_to_index = [
            Document(
                page_content=chunk,
                metadata={
                    "source": filename,
                    "chunk_id": i,
                    "file_type": file_extension
                }
            ) for i, chunk in enumerate(chunks)
        ]

        # Add documents to Azure AI Search
        vector_store.add_documents(docs_to_index)
        return True

    except HttpResponseError as e:
        if "Storage quota has been exceeded" in str(e):
            raise ValueError("Storage quota exceeded. Please delete documents or upgrade your Azure Search SKU.")
        else:
            raise

    except Exception as e:
        raise RuntimeError(f"An unexpected error occurred during RAG processing: {str(e)}")


def answer_question_from_docs(question):
    retrieved_docs = vector_store.similarity_search(query=question, k=3)
 
    if not retrieved_docs:
        return "No relevant information found in the indexed documents for your question."
 
    context_for_llm = "\n\n".join([doc.page_content for doc in retrieved_docs])
 
    messages = [
        {"role": "system", "content": "You are a helpful assistant. "
        "Use the context provided to answer the question. "
        "Answer the question based ONLY on the provided context. "
        "The information in the answer should NOT be outside the context. "
        "If the question is not related to the context, inform that politely. "
        "If the answer is not in the context, inform that with a simple message without using the word context. "
        "You may make calculations if needed. "
        "You have to answer with a polite greeting. "
        "You can answer with polite responses without using the word context if the user is not asking questions and is talking but that too should be based on the context. "
        "Don't give any suggestions or recommendations out of the context. "
        "You can reply on feedback."
        "If the question is not clear, ask for clarification. "},
        {"role": "user", "content": f"Context:\n{context_for_llm}\n\nQuestion: {question}"}
    ]
 
    response = chat_client.chat.completions.create(
        model=AZURE_OPENAI_CHAT_DEPLOYMENT_NAME,
        messages=messages,
        temperature=0.6,
        max_tokens=500
    )
 
    return response.choices[0].message.content.strip()